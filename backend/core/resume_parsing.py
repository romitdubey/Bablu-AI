from config.groq_api import parse_resume_with_groq
import PyPDF2
from docx import Document

def extract_text_from_docx(docx_path):
    lines = []
    doc = Document(docx_path)
    for paragraph in doc.paragraphs:
        lines.append(paragraph.text + "\n")
        
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    lines.append(cell_text)
    return ' '.join(lines)

def extract_text_from_pdf(pdf_path):
    text = ""
    with open(pdf_path, 'rb') as file:
        reader = PyPDF2.PdfReader(file)
        for page in reader.pages:
            text += page.extract_text()
    return text

